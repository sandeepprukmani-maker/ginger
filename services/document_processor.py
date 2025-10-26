import os
import requests
from PyPDF2 import PdfReader
from docx import Document
from PIL import Image
import pytesseract
import pandas as pd
import tempfile
import logging

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        pass
    
    def process_attachment(self, attachment_url, filename, auth):
        """Process document attachments and extract text"""
        try:
            file_extension = os.path.splitext(filename)[1].lower()
            
            response = requests.get(attachment_url, auth=auth)
            response.raise_for_status()
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
                tmp_file.write(response.content)
                tmp_path = tmp_file.name
            
            try:
                if file_extension == '.pdf':
                    text = self._extract_pdf_text(tmp_path)
                elif file_extension in ['.docx', '.doc']:
                    text = self._extract_docx_text(tmp_path)
                elif file_extension == '.txt':
                    text = self._extract_txt_text(tmp_path)
                elif file_extension in ['.xlsx', '.xls']:
                    text = self._extract_excel_text(tmp_path)
                elif file_extension in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff']:
                    text = self._extract_image_text(tmp_path)
                else:
                    text = f"[Unsupported file type: {file_extension}]"
                
                return text
            finally:
                if os.path.exists(tmp_path):
                    os.unlink(tmp_path)
                    
        except Exception as e:
            logger.error(f"Error processing {filename}: {str(e)}")
            return f"[Error processing {filename}: {str(e)}]"
    
    def _extract_pdf_text(self, file_path):
        """Extract text from PDF files"""
        try:
            reader = PdfReader(file_path)
            text_parts = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            return '\n\n'.join(text_parts)
        except Exception as e:
            return f"[Error extracting PDF: {str(e)}]"
    
    def _extract_docx_text(self, file_path):
        """Extract text from DOCX files"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return '\n\n'.join(text_parts)
        except Exception as e:
            return f"[Error extracting DOCX: {str(e)}]"
    
    def _extract_txt_text(self, file_path):
        """Extract text from TXT files"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            return f"[Error extracting TXT: {str(e)}]"
    
    def _extract_excel_text(self, file_path):
        """Extract text from Excel files (.xlsx, .xls) using pandas
        
        Uses xlrd (1.2.0) for legacy .xls files and openpyxl for modern .xlsx files
        """
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.xls':
                excel_file = pd.ExcelFile(file_path, engine='xlrd')
            else:
                excel_file = pd.ExcelFile(file_path, engine='openpyxl')
            
            text_parts = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                text_parts.append(f"\n=== Sheet: {sheet_name} ===\n")
                
                for _, row in df.iterrows():
                    row_values = [str(val) if pd.notna(val) else '' for val in row]
                    if any(row_values):
                        text_parts.append(' | '.join(row_values))
            
            return '\n'.join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting Excel file {file_path}: {str(e)}")
            return f"[Error extracting Excel: {str(e)}]"
    
    def _extract_image_text(self, file_path):
        """Extract text from images using OCR"""
        try:
            image = Image.open(file_path)
            
            text = pytesseract.image_to_string(image)
            
            if text.strip():
                return text.strip()
            else:
                return "[No text detected in image]"
        except Exception as e:
            logger.error(f"Error extracting text from image: {str(e)}")
            return f"[Error performing OCR: {str(e)}]"
